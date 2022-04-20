import matplotlib
matplotlib.use('Agg')
matplotlib.rc('font', family='DejaVu Sans')
matplotlib.rcParams['figure.dpi'] = 240
matplotlib.rcParams['font.size'] = 22
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches

from odoo import api, fields, models, _
from odoo.exceptions import ValidationError
    
 
class WordGenerator(models.Model):
    _inherit = "survey.survey"

    @api.multi
    def download_word(self):
        title = str(self.title)
        document = Document()
        document.add_heading(title, 1)
        id_survey = self.id
        survey_page = self.env['survey.page'].search([('survey_id', '=', self.id)])
        for page_svy in survey_page:
            document.add_heading(page_svy.title, level=1)
            question = self.env['survey.question'].search([('page_id', '=', page_svy.id)])
            for quest in question:
                document.add_heading(quest.question, level=1)
                if quest.type == 'simple_choice':
                    self.env.cr.execute("""
                                            SELECT a.value_suggested, COUNT(a.id) as total
                                            FROM survey_user_input_line a
                                            LEFT JOIN survey_user_input b on b.id = a.user_input_id
                                            WHERE a.question_id = %s and b.state = 'done'
                                            GROUP BY a.value_suggested
                                        """,(int(quest.id),))
                    dataSrv = self.env.cr.dictfetchall()
                    totalData = self.env['survey.user_input_line'].search_count([('question_id', '=', quest.id)])
                    
                    bars = []
                    height = []
                    for data in dataSrv:
                        nilai = str(data['value_suggested'])
                        if nilai is not None:
                            aa = data['value_suggested']
                            name = self.env['survey.label'].search([('id', '=', aa)])
                            nama = str(name.value)
                        else:
                            nama = str(data['value_suggested'])

                        bars.append(nama)
                        height.append(float(data['total']))
                    
                        def absolute_value(val):
                            a  = np.round(val/100.*np.array(height).sum(), 0)
                            res = int(a)
                            return res
                        
                        fig, ax = plt.subplots(figsize=(8,8), subplot_kw=dict(aspect="equal"))
                        theme = plt.get_cmap('hsv')
                        ax.set_prop_cycle("color", [theme(1. * i / len(bars)) for i in range(len(bars))])
                        wedges, texts, juck = ax.pie(height, startangle=-40, autopct=absolute_value)
                        kw = dict(arrowprops=dict(arrowstyle="-"), zorder=0, va="center")
                        for i, p in enumerate(wedges):
                            ang = (p.theta2 - p.theta1)/5. + p.theta1
                            y = np.sin(np.deg2rad(ang))
                            x = np.cos(np.deg2rad(ang))
                            horizontalalignment = {-1: "right", 1: "left"}[int(np.sign(x))]
                            connectionstyle = "angle,angleA=0,angleB={}".format(ang)
                            kw["arrowprops"].update({"connectionstyle": connectionstyle})
                            ax.annotate(bars[i], xy=(x, y), xytext=(1.35*np.sign(x), 1.4*y),
                                        horizontalalignment=horizontalalignment, **kw)

                    plt.savefig("/home/sippol/odoo/custom-addons/ms_download/static/doc/pie.png", bbox_inches = 'tight')
                    p = document.add_paragraph()
                    r = p.add_run()
                    pic = "/home/sippol/odoo/custom-addons/ms_download/static/doc/pie.png"
                    r.add_picture(pic, width=Inches(4),height=Inches(4))
                    plt.close()
        
                elif quest.type == 'textbox':
                    self.env.cr.execute("""
                                            SELECT a.value_text
                                            FROM survey_user_input_line a
                                            LEFT JOIN survey_user_input b on b.id = a.user_input_id
                                            WHERE a.question_id = %s AND 
                                            b.state = 'done' AND
                                            a.value_text IS NOT NULL
                                            GROUP BY value_text
                                        """,(int(quest.id),))
                    dataText             = self.env.cr.dictfetchall()
                    table               = document.add_table(rows=1, cols=2,style='Table Grid')
                    hdr_cells           = table.rows[0].cells
                    hdr_cells[0].text   = 'Input'
                    hdr_cells[1].text   = 'Responden'
                    cnt = 1
                    for line_tabel in dataText:
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(cnt)
                        row_cells[1].text = str(line_tabel['value_text'].encode('ascii','ignore').decode('ascii'))
                        cnt = cnt + 1

                elif quest.type == 'free_text':
                    self.env.cr.execute("""
                                            SELECT a.value_free_text
                                            FROM survey_user_input_line a
                                            LEFT JOIN survey_user_input b on b.id = a.user_input_id
                                            WHERE a.question_id = %s AND
                                            b.state = 'done' AND
                                            a.value_free_text IS NOT NULL
                                            GROUP BY a.value_free_text
                                        """,(int(quest.id),))
                    user_line2 = self.env.cr.dictfetchall()              
                    table = document.add_table(rows=1, cols=2,style='Table Grid')
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Input'
                    hdr_cells[1].text = 'Responden'
                    cnt = 1
                    for tab2 in user_line2:
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(cnt)
                        row_cells[1].text = str(tab2['value_free_text'].encode('ascii','ignore').decode('ascii'))
                        cnt = cnt + 1

                elif quest.type == 'numerical_box':
                    self.env.cr.execute("""
                                            SELECT a.value_number, AVG(a.value_number),COUNT(a.id) as total
                                            FROM survey_user_input_line a
                                            LEFT JOIN survey_user_input b on b.id = a.user_input_id
                                            WHERE a.question_id = %s AND 
                                            b.state = 'done' AND
                                            a.value_number IS NOT NULL
                                            GROUP BY a.value_number
                                        """,(int(quest.id),))
                    num_box  = self.env.cr.dictfetchall()
                    table = document.add_table(rows=1, cols=2,style='Table Grid')
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Input'
                    hdr_cells[1].text = 'Responden'
                    cnt = 1
                    for tab2 in num_box:
                        row_cells           = table.add_row().cells
                        row_cells[0].text   = str(cnt)
                        row_cells[1].text   = str(tab2['value_number'])
                        cnt = cnt + 1

                elif quest.type == 'datetime':
                    self.env.cr.execute("""
                                            SELECT a.value_date, COUNT(a.id) as total
                                            FROM survey_user_input_line a
                                            LEFT JOIN survey_user_input b on b.id = a.user_input_id
                                            WHERE a.question_id = %s AND
                                            b.state = 'done' AND
                                            a.value_date IS NOT NULL
                                            GROUP BY a.value_date
                                        """,(int(quest.id),))
                    date  = self.env.cr.dictfetchall()
                    table = document.add_table(rows=1, cols=2,style='Table Grid')
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Input'
                    hdr_cells[1].text = 'Responden'
                    cnt = 1
                    for x in date:
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(cnt)
                        row_cells[1].text = str(x['value_date'])
                        cnt = cnt+1

                elif quest.type == 'multiple_choice':
                    totalData = self.env['survey.user_input_line'].search_count([('question_id', '=', quest.id)])
                    table = document.add_table(rows=1, cols=3,style='Table Grid')
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Input'
                    hdr_cells[1].text = 'Responden'
                    hdr_cells[2].text = 'Persentase(%)'

                    self.env.cr.execute("""
                                            SELECT a.value_suggested, COUNT(a.id) as total
                                            FROM survey_user_input_line a
                                            LEFT JOIN survey_user_input b on b.id = a.user_input_id
                                            WHERE a.question_id = %s AND 
                                            b.state = 'done' AND
                                            a.value_suggested IS NOT NULL
                                            GROUP BY a.value_suggested
                                        """,(int(quest.id),))
                    multi_choice = self.env.cr.dictfetchall()
                    
                    bars = []
                    height = []
                    for data in multi_choice:
                        persentase = (float(data['total']) / float(totalData))*100
                        persen = round(persentase, 2)
                        nilai = str(data['value_suggested'])
                        if nilai is not None:
                            aa = data['value_suggested']
                            name = self.env['survey.label'].search([('id', '=', aa)])
                            nama = str(name.value)
                        else:
                            nama = str(data['value_suggested'])
                        
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(nama)
                        row_cells[1].text = str(data['total'])
                        row_cells[2].text = str(persen)+'%'
                        bars.append(nama)
                        height.append(float(data['total']))

                elif quest.type == 'matrix':
                    self.env.cr.execute("""
                                            SELECT DISTINCT a.value_suggested_row as id, b."value" as value
                                            FROM survey_user_input_line a
                                            LEFT JOIN survey_label b on a.value_suggested_row = b.id
                                            LEFT JOIN survey_user_input c on c.id = a.user_input_id
                                            WHERE a.question_id = %s and 
                                            c.state = 'done' AND
                                            a.value_suggested_row IS NOT NULL
                                            ORDER BY a.value_suggested_row
                                        """,(int(quest.id),))
                    pertanyaan = self.env.cr.dictfetchall()
                    
                    idx_total = 0
                    for row in pertanyaan:
                        self.env.cr.execute("""
                                            SELECT c.id, a.value_suggested_row, b."value", c."value" as jawaban, COUNT(c.id) as total
                                            FROM survey_user_input_line a
                                            LEFT JOIN survey_label b on a.value_suggested_row = b.id
                                            LEFT JOIN survey_label c on a.value_suggested = c.id
                                            LEFT JOIN survey_user_input d on d.id = a.user_input_id
                                            WHERE d.state = 'done' AND
                                            a.value_suggested_row = %s
                                            GROUP BY c."id", a.value_suggested_row, b."value"
                                        """,(int(row['id']),))
                        data = self.env.cr.dictfetchall()

                        bars = [] # Jawaban
                        height = [] # Total
                        for x in data:
                            bars.append(str(x['jawaban']))
                            height.append(int(x['total']))
                        
                        cols = len(bars)
                        table = document.add_table(rows=1, cols=cols+1, style='Table Grid')
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text = 'Pertanyaan'
                        hc = 1
                        bars_idx = 0
                        for x in range(cols):
                            hdr_cells[hc].text = str(bars[bars_idx])
                            hc = hc + 1
                            bars_idx = bars_idx + 1

                        row_cells = table.add_row().cells
                        row_cells[0].text = row['value'] 

                        row_idx = 1
                        height_idx = 0
                        for x in range(cols):
                            row_cells[row_idx].text = str(height[height_idx])
                            row_idx = row_idx + 1
                            height_idx = height_idx + 1
                            
                        idx_total = idx_total + 1
                        
                        p = document.add_paragraph()
                        p.add_run()

        document.save("/home/sippol/odoo/custom-addons/ms_download/static/doc/HasilSurvey.docx")

        return {
            'type': 'ir.actions.act_url',
            'url': '/ms_download/static/doc/HasilSurvey.docx',            
            'target': 'new',
        }

 